from flask import Flask, request, jsonify, send_from_directory, send_file , render_template
import os
import requests
from dotenv import load_dotenv
import psycopg2
from flask_cors import CORS





from io import BytesIO


from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT



load_dotenv()
app = Flask(__name__)
CORS(app,origins=["https://e45m.github.io","https://pyred.com.co"])

API_KEY = os.getenv('GEMINI_API_KEY')
GEMINI_URL = f'https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={API_KEY}'


##############################


DATABASE_URL = os.getenv("DATABASE_URL")
def get_connection():
    return psycopg2.connect(DATABASE_URL)

@app.route('/vX3zqMnAyB7gRtCH9fLdNpJE0WkTuhQ2xVaPiLsoY4cbUGjKeTwZ2', methods=['GET'])
def obtener_visitas():
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT contador FROM visitas LIMIT 1;")
            visitas = cur.fetchone()[0]
    return {'visitas': visitas}, 200


@app.route('/vX3zqMnAyB7gRtCH9fLdNpJE0WkTuhQ2xVaPiLsoY4cbUGjKeTwZ', methods=['POST'])
def reset_visitas():
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("DROP TABLE IF EXISTS visitas;")
            cur.execute("""
                CREATE TABLE visitas (
                    id SERIAL PRIMARY KEY,
                    contador INTEGER NOT NULL
                );
            """)
            cur.execute("INSERT INTO visitas (contador) VALUES (0);")
        conn.commit()
    return "Base de datos reiniciada."



def init_db():
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                CREATE TABLE IF NOT EXISTS visitas (
                    id SERIAL PRIMARY KEY,
                    contador INTEGER NOT NULL
                );
            """)
            cur.execute("SELECT COUNT(*) FROM visitas;")
            if cur.fetchone()[0] == 0:
                cur.execute("INSERT INTO visitas (contador) VALUES (0);")
        conn.commit()

def incrementar_visitas():
    with get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("UPDATE visitas SET contador = contador + 1 RETURNING contador;")
            visitas = cur.fetchone()[0]
        conn.commit()
    return visitas





#############################



texto_global = ''

@app.route('/wakeup', methods=['GET'])
def wakeup():


        return jsonify({"respuesta": "Hola, Soy la API del servicio, ya estoy lista para responder....\npuede que tarde en responder por favor ten paciencian unos minutos"})

@app.route("/")
def redirect_page():
    return render_template("redirect.html")  # Renderiza el archivo redirect.html





@app.route('/consulta', methods=['POST'])
def consulta():
    incrementar_visitas()
    global texto_global

    data = request.get_json()
    cargo = data['cargo']
    empresa = data['empresa']
    ciudad = data['ciudad']
    cv = data['cv']


    prompt =  f"""Este es mi currículum (hoja de vida):\n\n{cv}\n\nTenlo en cuenta para responder las siguientes solicitudes:\n 1: Análisis del CV para el puesto de {cargo}. Revisa mi currículum actual en relación con el puesto indicado. Identifica oportunidades de mejora, incluyendo: palabras clave relevantes del sector que faltan, logros que deben cuantificarse, habilidades técnicas ausentes y posibles mejoras en la estructura general. Compara mi perfil con la descripción del puesto y sugiere cambios concretos. Prioriza los elementos de alto impacto y elimina contenido irrelevante. Proporciona ejemplos de frases mejoradas con verbos de acción fuertes y métricas específicas.\n\n 2: Preparación de historias STAR para entrevistas. Basado en mi experiencia, ayúdame a construir tres historias siguiendo el método STAR: un proyecto exitoso, un conflicto que resolví y un fracaso que superé. Para cada historia, estructura claramente la SITUACIÓN, la TAREA, las ACCIONES y los RESULTADOS obtenidos. Incluye frases de transición efectivas y momentos clave para asegurar que cada historia no exceda los 2 minutos de duración.\n\n 3: Análisis estratégico de la empresa {empresa}. Investiga y resume la posición actual de esta empresa en el mercado. Destaca sus valores fundamentales, logros recientes, desafíos del sector y noticias relevantes de los últimos 6 meses. Sugiere aspectos clave que pueda mencionar en una entrevista para demostrar interés genuino. Conecta mi perfil con la industria, el puesto y los valores de la empresa.\n\n 4: Simulación de entrevista técnica para el puesto de {cargo}. Crea una entrevista técnica simulada con preguntas específicas del sector, casos prácticos y situaciones complejas que podría enfrentar en el cargo. Evalúa mis posibles respuestas con base en mi CV y señala áreas a mejorar. Identifica brechas potenciales en mi perfil y sugiere cómo cubrirlas.\n\n 5: Diferenciación estratégica para destacar como candidato. ¿Cómo puedo posicionarme de manera única para este puesto en {empresa}? Analiza las tendencias del mercado y sugiere enfoques innovadores para presentar mi experiencia. Identifica factores diferenciadores frente al candidato promedio y cómo comunicar esas fortalezas efectivamente.\n\n 6: Estrategia de negociación salarial en {ciudad}. ¿Cuál es el rango salarial competitivo actual para el puesto de {cargo} en {ciudad }? Diseña una estrategia de negociación que incluya beneficios no monetarios y puntos clave de inflexión para maximizar la propuesta total de valor. Genera solo texto limpio y sin marcas de formato. Responde en el orden indicado. Encabeza cada sección con un título profesional. Limita cada respuesta a un máximo de 512 palabras. Separa cada respuesta con tres saltos de línea. Redacta como una firma consultora de alto nivel, especializada en imagen profesional y recursos humanos.Analyze the provided CV. Analiza el CV. Si no está en español, indica el idioma en el que está escrito y responde en ese idioma. """





    payload = {
        "contents": [
            {
                "parts": [
                    {
                        "text": prompt
                    }
                ]
            }
        ]
    }

    try:
        response = requests.post(GEMINI_URL, json=payload)
        response.raise_for_status()

        resultado = response.json()
        texto = resultado.get("candidates", [{}])[0].get("content", {}).get("parts", [{}])[0].get("text", "Sin respuesta.")
        texto = texto.replace('*','')
        texto_global = f'\n {texto} \n \n \n'

        return jsonify({"respuesta": texto})

    except Exception as e:
        return jsonify({"error": str(e)+'***\n'+str(response.text)}), 500




@app.route('/download-word')
def download_word():

    # Crear un documento de Word
    doc = Document()

    # Cambiar la orientación a horizontal
    # section = doc.sections[-1]
    # section.orientation = WD_ORIENT.LANDSCAPE

    # Ajustar el tamaño de la página al nuevo ancho
    # new_width, new_height = section.page_height, section.page_width
    # section.page_width = new_width
    # section.page_height = new_height
    global texto_global

    doc.add_heading('Informe de Optimización de CV con IA', 0)

    doc.add_paragraph(texto_global)







    # Guardar el documento en memoria
    output = BytesIO()
    doc.save(output)
    output.seek(0)

    # Enviar el archivo Word como respuesta
    return send_file(output, download_name="recomendacionesCV.docx", as_attachment=True)


if __name__ == '__main__':
    app.run(debug=True)
